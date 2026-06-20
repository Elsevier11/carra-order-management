from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"C:\Dev\Carra_Ordini\Istruzioni_deploy_server_cliente.docx")


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.1875)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.1875)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(value)
    set_font(r2)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["Normal"].font.size = Pt(11)

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title_run = title.add_run("Istruzioni di deploy server cliente")
set_font(title_run, size=26, bold=False, color="000000")

sub = doc.add_paragraph()
set_paragraph_spacing(sub, before=0, after=8)
sub_run = sub.add_run(
    "Pacchetto operativo per installazione Linux in Docker, con import iniziale dati e collaudo base."
)
set_font(sub_run, size=11, color="555555")

doc.add_heading("Contesto", level=1)
p = doc.add_paragraph()
set_paragraph_spacing(p, after=6)
run = p.add_run(
    "Questa procedura e pensata per il fornitore che installa l'applicativo sul server Linux del cliente. "
    "L'applicazione viene eseguita in Docker, con frontend, backend e database nello stesso stack."
)
set_font(run)

doc.add_heading("Repository e branch", level=2)
add_bullet(doc, "Repository: https://github.com/Elsevier11/carra-order-management.git")
add_bullet(doc, "Branch da utilizzare: main")
add_bullet(doc, "Il repository contiene solo il codice applicativo, non dati cliente pre-caricati.")

doc.add_heading("Stack previsto", level=2)
add_bullet(doc, "Frontend Angular servito da Nginx")
add_bullet(doc, "Backend Node.js / Express")
add_bullet(doc, "Database PostgreSQL")
add_bullet(doc, "Frontend e API esposti sullo stesso origin per evitare problemi di CORS")

doc.add_heading("Prerequisiti", level=2)
add_bullet(doc, "Docker Engine")
add_bullet(doc, "Docker Compose")
add_bullet(doc, "Accesso al repository GitHub")
add_bullet(doc, "Spazio persistente sul server per database e allegati caricati dall'app")

doc.add_heading("Variabili minime", level=2)
add_bullet(doc, "POSTGRES_DB")
add_bullet(doc, "POSTGRES_USER")
add_bullet(doc, "POSTGRES_PASSWORD")
add_bullet(doc, "JWT_SECRET")

doc.add_heading("Avvio dello stack", level=2)
p = doc.add_paragraph()
set_paragraph_spacing(p, after=4)
r = p.add_run("Comando di avvio: ")
set_font(r, bold=True)
r = p.add_run("docker compose up -d --build")
set_font(r)

doc.add_heading("Import iniziale dati", level=2)
add_bullet(doc, "Il database va popolato una sola volta prima del primo avvio operativo.")
add_bullet(doc, "L'import non avviene dall'interfaccia web, ma tramite lo script previsto nel repository.")
add_bullet(doc, "Il file da usare e il JSON iniziale gia predisposto, con lo stesso contenuto della mia cartella.")
add_bullet(doc, "In questo modo il cliente parte con il medesimo set di dati gia pronto per il test.")

p = doc.add_paragraph()
set_paragraph_spacing(p, after=6)
r = p.add_run("Comando di riferimento: ")
set_font(r, bold=True)
r = p.add_run("npm run db:import -- ./data/consegne.full.json")
set_font(r)

doc.add_heading("Verifiche richieste", level=2)
add_number(doc, "GET /health")
add_number(doc, "Login applicativo")
add_number(doc, "Visualizzazione elenco ordini")
add_number(doc, "Modifica ordine")
add_number(doc, "Eventuale upload allegati, se incluso nel collaudo")

doc.add_heading("Nota operativa", level=2)
p = doc.add_paragraph()
set_paragraph_spacing(p, after=6)
run = p.add_run(
    "In ambiente di test, se necessario, l'import puo essere ripetuto. "
    "In produzione, eventuali reimport completi vanno fatti solo con backup e conferma preventiva."
)
set_font(run)

doc.add_heading("Riscontro richiesto", level=2)
add_number(doc, "URL finale di pubblicazione")
add_number(doc, "Esito dell'health check")
add_number(doc, "Conferma del collaudo base")

doc.save(OUT_PATH)
print(OUT_PATH)
